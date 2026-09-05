from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/silverbrick/Documents/ChatGPT/arm_eye/arm_eye答辩逐页讲稿.docx"

# Preset: compact_reference_guide. Named override: Heiti SC bundled for render QA.
FONT = "Heiti SC"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_FILL = "F4F6F9"
BLUE_FILL = "E8EEF5"
GOLD = "7A5A00"
BORDER = "D5DDE7"

slides = [
    {
        "title": "视觉手势驱动的四自由度机械臂控制系统",
        "time": "30–35 秒",
        "keywords": "项目定位｜arm + arm_eye｜非接触控制",
        "script": "各位老师好，我答辩的题目是《视觉手势驱动的四自由度机械臂控制系统》。本项目包含 arm 和 arm_eye 两部分：arm 负责机械臂的基础控制、动作录制与回放；arm_eye 在此基础上加入电脑摄像头手势识别，实现更加自然的非接触控制。下面我将介绍项目演进、系统实现、安全机制、应用前景和现场演示。",
        "bridge": "先说明两个项目之间的关系，再进入具体实现。",
    },
    {
        "title": "从示教复现到自然交互",
        "time": "50–55 秒",
        "keywords": "示教复现｜继承关系｜控制入口升级",
        "script": "arm 是基础项目，通过四个电位器控制四个舵机，并支持动作录制和回放，解决的是机械臂如何稳定、安全地运动。arm_eye 没有修改 arm 项目，而是在其可靠控制基础上增加摄像头、手势识别和串口命令层，把输入方式从机械电位器扩展成自然手势。因此两者是继承关系：先完成可靠控制，再实现自然交互。",
        "bridge": "有了这个分工，下面先看 arm 如何建立可靠的四轴控制底座。",
    },
    {
        "title": "arm 项目：建立可靠的四轴控制底座",
        "time": "50 秒",
        "keywords": "四轴控制｜ADC 标定｜录制回放",
        "script": "arm 项目首先建立四轴控制底座。四个关节分别读取电位器信号，并通过 ADC 范围标定把输入映射成对应关节角度。GPIO13 按钮支持动作录制和回放，录制保存的是完整运动过程，而不仅是起始和结束角度，因此能够复现连续轨迹。安全保护机制放在后面的独立章节中说明。",
        "bridge": "在稳定的固件底座上，再叠加视觉识别和串口通信，就形成了 arm_eye。",
    },
    {
        "title": "arm_eye：构建“看见—判断—执行”闭环",
        "time": "60–65 秒",
        "keywords": "摄像头｜MediaPipe｜串口｜ESP32",
        "script": "arm_eye 形成一条完整链路：首先调用电脑或外接摄像头采集画面，然后由 MediaPipe 输出手势类别和置信度；Python 程序经过稳定性判断后，将手势映射成单字符命令，通过 USB 串口发送给 ESP32；最后 ESP32 把指令转换成安全、平滑的舵机运动。识别与固件分层，摄像头、串口和舵机可以独立调试。",
        "bridge": "接下来把软件链路落到硬件引脚和关节角度上。",
    },
    {
        "title": "硬件映射与关节校准",
        "time": "50–55 秒",
        "keywords": "J1/J2/J3/J4｜PWM｜安全角度｜J2=78°",
        "script": "四个关节使用独立的电位器和 PWM 输出：J1 控制底座旋转，J2 用于稳定整体姿态，J3 负责抬升，J4 负责末端夹爪。安全范围沿用了 arm 项目的校准结果。抓取动作中，J2 固定在 78 度，这是实际调试得到的稳定值，可以减少 J3 和 J4 运动时的姿态漂移。",
        "bridge": "完成硬件校准后，就可以把可识别的手势逐一映射成动作。",
    },
    {
        "title": "五种手势映射为可解释动作",
        "time": "65–70 秒",
        "keywords": "S｜G｜U/D｜R｜动作解释｜安全角度",
        "script": "系统支持五种手势。Open_Palm 对应 S，表示待机或松开，J4 调整到最大安全角；Closed_Fist 对应 G，触发录制的抓取过程，其中 J4 完成爪夹闭合，J3 配合完成姿态变化；Pointing_Up 对应 U，J3 相对当前位置逆时针 15 度；Thumb_Down 对应 D，是 U 的反动作；Victory 对应 R，控制底座 J1 旋转。每个手势都能对应到具体关节和明确动作，便于校准和演示。另外，所有动作的目标角度都受 arm 项目校准的安全角度范围限制，四个关节不会越过各自的安全边界；安全角度负责兜底，手势映射负责定义动作。",
        "bridge": "手势映射解决的是“做什么”，下面说明系统如何避免误识别和过冲。",
    },
    {
        "title": "安全性：识别过滤 + 运动约束",
        "time": "60–65 秒",
        "keywords": "置信度 0.65｜连续 4 帧｜2°/20 ms｜复位",
        "script": "视觉识别可能出现瞬时误判，因此系统采用多层安全保护。置信度必须达到 0.65 才进入候选状态；同一手势连续稳定 4 帧后才发送指令；进入执行阶段后仍受安全角度和每 20 ms 最大变化 2 度的限制。如果出现异常，还可以长按 GPIO13 两秒，让四个关节平滑回到 90 度。安全性不是只依赖识别程序，而是从视觉、通信到舵机执行进行了多层限制。",
        "bridge": "在安全边界明确之后，系统具备从原型走向实际应用的基础。",
    },
    {
        "title": "应用前景：从手势控制走向人机协作",
        "time": "40 秒",
        "keywords": "搬运分拣｜危险环境｜康复辅助｜教学平台",
        "script": "在现有原型基础上，系统可以向四类场景扩展。第一是轻型物料的搬运和分拣；第二是在高温、粉尘等危险环境中，让操作者与机械臂保持距离；第三是康复训练和无接触辅助服务；第四是作为教学平台，验证视觉识别、串口通信和机器人控制算法。需要说明的是，这些属于后续应用方向，当前项目仍处于功能原型阶段。",
        "bridge": "下面用现场演示说明这条链路如何运行，并交代当前还可以怎样优化。",
    },
    {
        "title": "现场演示路径与后续优化",
        "time": "60–65 秒",
        "keywords": "烧录｜识别窗口｜S/U/D/R/G｜复位",
        "script": "现场演示建议按从低风险到高风险的顺序进行：先双击启动器并完成烧录，再确认摄像头显示识别标签，然后依次演示 S、U/D、R、G，最后长按按钮复位。当前主要限制不在识别链路，而在夹爪结构、摩擦力和抓取轨迹时长。后续可以从防滑夹爪、轨迹压缩、目标检测和力反馈等方向优化。",
        "bridge": "最后用一句话概括两个项目各自解决的问题和当前成果。",
        "demo_note": "若现场摄像头或串口出现异常，不要反复调试；可直接说明识别窗口、串口日志或预录视频作为兜底，并继续讲解控制链路。",
    },
    {
        "title": "总结",
        "time": "30–35 秒",
        "keywords": "可靠控制底座｜视觉闭环｜可扩展原型",
        "script": "arm 项目完成了机械臂可靠控制的底座；arm_eye 在此基础上实现了从视觉识别到安全执行的完整闭环。项目目前已经具备可运行、可测试和可继续扩展的工程基础。我的汇报结束，谢谢各位老师，请批评指正。",
        "bridge": "—",
    },
]

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn('w:' + name))
        if node is None:
            node = OxmlElement('w:' + name)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_table_borders(table, color=BORDER, size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.find(qn('w:' + edge))
        if element is None:
            element = OxmlElement('w:' + edge)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)

def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(widths_dxa[index]))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn('w:ascii'), FONT)
    r_pr.rFonts.set(qn('w:hAnsi'), FONT)
    r_pr.rFonts.set(qn('w:eastAsia'), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_pfmt(paragraph, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True

def add_field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instr_text)
    run._r.append(end)
    set_run(run, size=9, color=MUTED)

def add_label_detail_table(doc, rows, fill=LIGHT_FILL):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], BLUE_FILL)
        set_cell_shading(cells[1], fill)
        p0 = cells[0].paragraphs[0]
        set_pfmt(p0, before=0, after=0, line=1.15)
        set_run(p0.add_run(label), size=9.5, color=DARK_BLUE, bold=True)
        p1 = cells[1].paragraphs[0]
        set_pfmt(p1, before=0, after=0, line=1.15)
        set_run(p1.add_run(value), size=9.5, color=INK)
    set_table_geometry(table, [1700, 7660])
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table

def add_callout(doc, title, text, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_table_geometry(table, [9360])
    set_table_borders(table, color='C9D5E4', size='8')
    p = cell.paragraphs[0]
    set_pfmt(p, before=0, after=3, line=1.15, keep=True)
    set_run(p.add_run(title), size=10, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_pfmt(p2, before=0, after=0, line=1.25)
    set_run(p2.add_run(text), size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table

def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style='List Bullet')
    set_pfmt(paragraph, before=0, after=4, line=1.25)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    set_run(paragraph.add_run(text), size=10.5, color=INK)
    return paragraph

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn('w:ascii'), FONT)
normal._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 18, 10),
    ('Heading 2', 13, BLUE, 14, 7),
    ('Heading 3', 12, DARK_BLUE, 10, 5),
]:
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn('w:ascii'), FONT)
    style._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.keep_with_next = True

for style_name in ['List Bullet', 'List Number']:
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn('w:ascii'), FONT)
    style._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

# Quiet running header and footer.
header = section.header
header_p = header.paragraphs[0]
set_pfmt(header_p, before=0, after=0, line=1.0)
set_run(header_p.add_run('arm / arm_eye  ·  答辩逐页讲稿'), size=9, color=MUTED, bold=True)
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_pfmt(footer_p, before=0, after=0, line=1.0)
set_run(footer_p.add_run('答辩参考稿  ·  第 '), size=9, color=MUTED)
add_field(footer_p, 'PAGE')
set_run(footer_p.add_run(' 页'), size=9, color=MUTED)

# Cover / usage block. Pattern: editorial_cover with compact whitespace override.
p = doc.add_paragraph()
set_pfmt(p, before=26, after=6, line=1.0)
set_run(p.add_run('答辩演讲参考稿'), size=11, color=GOLD, bold=True)
p = doc.add_paragraph()
set_pfmt(p, before=0, after=8, line=1.05)
set_run(p.add_run('arm 与 arm_eye 项目答辩逐页讲稿'), size=26, color=INK, bold=True)
p = doc.add_paragraph()
set_pfmt(p, before=0, after=18, line=1.15)
set_run(p.add_run('10 分钟演讲参考稿｜答辩人：________｜2026 年 8 月'), size=12.5, color=MUTED)
add_callout(doc, '使用方法', '这份讲稿与精简版 PPT 的 10 页一一对应。建议不要逐字背诵，而是记住每页的“核心结论—证据—衔接”。按建议时长讲完约 9 分钟，剩余时间留给现场演示和提问。', fill=BLUE_FILL)
p = doc.add_paragraph(style='Heading 2')
p.add_run('全篇结构')
for run in p.runs:
    set_run(run, size=13, color=BLUE, bold=True)
add_label_detail_table(doc, [
    ('01–03', '项目定位、arm 控制底座与动作录制'),
    ('04–06', '视觉识别闭环、硬件校准与手势映射'),
    ('07–10', '安全机制、应用前景、演示路径与总结'),
])

doc.add_page_break()

for index, slide in enumerate(slides, start=1):
    heading = doc.add_paragraph(style='Heading 1')
    heading.add_run(f'第 {index:02d} 页｜{slide["title"]}')
    for run in heading.runs:
        set_run(run, size=16, color=BLUE, bold=True)
    add_label_detail_table(doc, [('建议时长', slide['time']), ('关键词', slide['keywords'])])
    add_callout(doc, '讲解稿', slide['script'])
    if slide.get('demo_note'):
        add_callout(doc, '现场提示', slide['demo_note'], fill='FFF8E6')
    p = doc.add_paragraph()
    set_pfmt(p, before=1, after=8, line=1.15)
    set_run(p.add_run('页间衔接：'), size=10, color=GOLD, bold=True)
    set_run(p.add_run(slide['bridge']), size=10, color=MUTED, italic=True)

doc.add_page_break()
heading = doc.add_paragraph(style='Heading 1')
heading.add_run('答辩速记：三句话与常见追问')
for run in heading.runs:
    set_run(run, size=16, color=BLUE, bold=True)

heading = doc.add_paragraph(style='Heading 2')
heading.add_run('可主动强调的三句话')
for run in heading.runs:
    set_run(run, size=13, color=BLUE, bold=True)
for text in [
    'arm 解决的是机械臂如何稳定运动，arm_eye 解决的是人如何自然控制机械臂。',
    '安全性由视觉过滤、舵机限速、角度范围和硬件复位共同保证。',
    '当前是功能原型，抓取能力和识别准确率还需要进一步定量测试。',
]:
    add_bullet(doc, text)

heading = doc.add_paragraph(style='Heading 2')
heading.add_run('常见追问速答')
for run in heading.runs:
    set_run(run, size=13, color=BLUE, bold=True)
qa = [
    ('为什么把识别放在电脑端？', 'MediaPipe 在电脑端运行更容易调试，也能减轻 ESP32 的计算负担；代价是系统依赖电脑和摄像头。'),
    ('0.65 和连续 4 帧是怎么确定的？', '这是在调试中对准确率与响应速度做的折中：门槛太低容易误触发，太高又会影响可操作性。目前还没有建立大规模标注数据集。'),
    ('为什么 G 不能只发送最终角度？', '抓取过程包含“先打开、再闭合”的时序，起点和终点角度无法表达中间极点，所以 G 使用录制的完整轨迹。'),
    ('为什么把 J2 固定在 78 度？', '受到重心和舵机性能限制，J2 在大负载时难以自由移动。多次测试表明 78 度能保持较稳定的姿态。'),
    ('抓不住物体时如何改进？', '问题可能来自夹爪摩擦力、舵机扭矩、物体形状或轨迹时序；后续可增加防滑结构、优化角度和速度，并加入力反馈。'),
    ('识别误判时如何保护？', '置信度、连续帧确认、角度范围和运动限速共同降低风险；异常时长按 GPIO13 两秒，让关节平滑回到 90 度。'),
]
for question, answer in qa:
    h = doc.add_paragraph(style='Heading 3')
    h.add_run(question)
    for run in h.runs:
        set_run(run, size=12, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    set_pfmt(p, before=0, after=7, line=1.25)
    set_run(p.add_run(answer), size=10.5, color=INK)

p = doc.add_paragraph()
set_pfmt(p, before=10, after=0, line=1.0)
set_run(p.add_run('资料依据：精简版答辩 PPT、arm_eye 项目源代码及手势控制脚本。'), size=9, color=MUTED, italic=True)

doc.core_properties.title = 'arm 与 arm_eye 项目答辩逐页讲稿'
doc.core_properties.subject = '10分钟答辩演讲参考稿'
doc.core_properties.author = 'Codex'
doc.core_properties.keywords = 'arm, arm_eye, 机械臂, 手势识别, 答辩'
doc.save(OUT)
print(OUT)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

path = "/Users/silverbrick/Documents/ChatGPT/arm_eye/arm_eye答辩逐页讲稿.docx"
doc = Document(path)

def set_run_font(run, size=10.5, color="0B2545"):
    run.font.name = "Heiti SC"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Heiti SC")
    rpr.rFonts.set(qn("w:hAnsi"), "Heiti SC")
    rpr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

# Slide 06 is the 7th metadata table and following callout table in the document.
metadata = doc.tables[12]
keyword_cell = metadata.cell(1, 1)
keyword_cell.paragraphs[0].add_run("｜安全角度")
set_run_font(keyword_cell.paragraphs[0].runs[-1])

script_cell = doc.tables[13].cell(0, 0)
script_paragraph = script_cell.paragraphs[1]
addition = "另外，所有动作的目标角度都受 arm 项目校准的安全角度范围限制，四个关节不会越过各自的安全边界；安全角度负责兜底，手势映射负责定义动作。"
run = script_paragraph.add_run(addition)
set_run_font(run)

doc.save(path)
print(path)
